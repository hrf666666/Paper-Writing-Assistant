import os
import re
import sys
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import latex2mathml.converter
import mathml2omml
from lxml import etree

# 确保XML命名空间正确设置
nsmap = {}
nsmap['m'] = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
nsmap['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

class MarkdownToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self._set_default_styles()
    
    def _set_default_styles(self):
        """设置文档中所有样式使用黑体字体。"""
        # 设置默认字体为黑体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'SimHei'  # 黑体
        font.size = Pt(11)
        
        # 将黑体应用到所有可用样式
        for style in self.doc.styles:
            if hasattr(style, 'font'):
                style.font.name = 'SimHei'  # 黑体
        
        # 对于标题，明确设置黑体
        for i in range(1, 10):  # 标题级别1-9
            if f'Heading {i}' in self.doc.styles:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = 'SimHei'  # 黑体
                heading_style.font.size = Pt(14 + (9-i))  # 标题使用更大的字体大小
                heading_style.font.bold = True  # 标题加粗
        
        # 设置列表的字体
        for list_style in ['List Bullet', 'List Number', 'List Bullet 2', 'List Number 2']:
            if list_style in self.doc.styles:
                self.doc.styles[list_style].font.name = 'SimHei'
        
        # 设置默认段落字体
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'SimHei'
    
    def _create_element(self, name):
        """创建具有给定名称的XML元素。"""
        return OxmlElement(name)
    
    def _create_attribute(self, element, name, value):
        """为给定元素创建属性。"""
        element.set(qn(name), value)
    
    def _latex_to_omml(self, latex_string):
        """将LaTeX转换为OMML格式。"""
        try:
            # 清理LaTeX公式
            latex_string = re.sub(r'\s+', ' ', latex_string).strip()
            
            # 第一步：LaTeX → MathML
            mathml = latex2mathml.converter.convert(latex_string)
            
            # 第二步：MathML → OMML
            omml = mathml2omml.convert(mathml)
            
            return omml
        except Exception as e:
            print(f"将LaTeX转换为OMML时出错: {e}", latex_string)
            return None

    def _insert_omml_into_paragraph(self, paragraph, omml_string):
        """将OMML插入到段落中。"""
        try:
            # 添加缺少的命名空间声明
            ns_prefixed_omml = omml_string.replace('<m:oMath>', '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">')
            
            run = paragraph.add_run()
            omml_element = parse_xml(ns_prefixed_omml)
            run._r.append(omml_element)
            return True
        except Exception as e:
            print(f"插入OMML时出错: {e}")
            print(f"OMML字符串内容: {omml_string[:100]}...")  # 打印部分OMML用于调试
            return False

    def _create_office_math_directly(self, latex):
        """直接创建Word数学公式对象作为备选方案。"""
        try:
            # 创建包含公式的oMathPara元素
            oMathPara = OxmlElement('m:oMathPara')
            oMath = OxmlElement('m:oMath')
            oMathPara.append(oMath)
            
            # 创建文本运行对象
            r = OxmlElement('m:r')
            t = OxmlElement('m:t')
            t.text = latex
            r.append(t)
            oMath.append(r)
            
            return oMathPara
        except Exception as e:
            print(f"直接创建公式对象时出错: {e}")
            return None
    
    def _process_formula(self, paragraph, formula_text):
        """处理LaTeX公式并将其插入段落。"""
        success = False
        
        # 规范化公式文本
        formula_text = formula_text.strip()
        
        # 方法1：使用mathml2omml库
        try:
            omml_string = self._latex_to_omml(formula_text)
            if omml_string is not None:
                if self._insert_omml_into_paragraph(paragraph, omml_string):
                    success = True
        except Exception as e:
            print(f"主要公式转换方法出错: {e}")
        
        # 方法2：如果方法1失败，则直接创建OMML结构
        if not success:
            try:
                math_element = self._create_office_math_directly(formula_text)
                if math_element is not None:
                    run = paragraph.add_run()
                    run._r.append(math_element)
                    success = True
            except Exception as e:
                print(f"备选公式转换方法出错: {e}")
        
        # 如果两种方法都失败，则添加原始文本
        if not success:
            paragraph.add_run(f"<formula>{formula_text}</formula>")
        
        return success
    
    def _add_page_break(self):
        """在文档中添加页面分隔符。"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _clean_markdown_text(self, text):
        """清除Markdown格式字符。"""
        # 移除Markdown格式字符
        cleaned_text = text
        # 移除用于格式化的星号（粗体，斜体）
        cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_text)  # 粗体
        cleaned_text = re.sub(r'\*(.*?)\*', r'\1', cleaned_text)      # 斜体
        
        # 根据需要移除其他Markdown语法
        cleaned_text = re.sub(r'`(.*?)`', r'\1', cleaned_text)        # 内联代码
        
        return cleaned_text
    
    def _format_code_block(self, code_text, language=""):
        """使用灰色背景和等宽字体格式化代码块。"""
        # 为代码块创建新段落
        p = self.doc.add_paragraph()
        
        # 为代码设置等宽字体，但保留黑体作为备选
        code_run = p.add_run(code_text)
        code_run.font.name = 'SimHei'  # 现在所有代码都使用黑体
        code_run.font.size = Pt(9)
        
        # 为段落添加灰色阴影
        shading_elm = self._create_element('w:shd')
        self._create_attribute(shading_elm, 'w:val', 'clear')
        self._create_attribute(shading_elm, 'w:color', 'auto')
        self._create_attribute(shading_elm, 'w:fill', 'E0E0E0')  # 浅灰色
        
        # 将阴影应用到段落
        p._p.get_or_add_pPr().append(shading_elm)
        
        # [注意] 以下代码已被注释或删除 - 2025-0319-23:33
        #
        # 语言标签显示
        # if language and language.strip():
        #     lang_p = self.doc.add_paragraph()
        #     lang_run = lang_p.add_run(f"语言: {language.strip()}")
        #     lang_run.font.name = 'SimHei'
        #     lang_run.font.size = Pt(8)
        #     lang_run.italic = True
    
    def _format_diagram(self, diagram_text):
        """使用等宽字体和浅色背景格式化ASCII艺术图。"""
        # 为图表创建新段落
        p = self.doc.add_paragraph()
        
        # 为图表设置等宽字体，但保留黑体作为主要字体
        diagram_run = p.add_run(diagram_text)
        diagram_run.font.name = 'SimHei'  # 现在所有图表都使用黑体
        diagram_run.font.size = Pt(9)
        
        # 为段落添加浅灰色阴影
        shading_elm = self._create_element('w:shd')
        self._create_attribute(shading_elm, 'w:val', 'clear')
        self._create_attribute(shading_elm, 'w:color', 'auto')
        self._create_attribute(shading_elm, 'w:fill', 'F5F5F5')  # 非常浅的灰色
        
        # 将阴影应用到段落
        p._p.get_or_add_pPr().append(shading_elm)
    
    def _process_heading(self, elem, level):
        """处理标题元素。"""
        heading_text = elem.get_text().strip()
        # 清除任何Markdown格式字符
        heading_text = self._clean_markdown_text(heading_text)
        
        heading = self.doc.add_heading(heading_text, level=level)
        # 确保标题使用黑体
        for run in heading.runs:
            run.font.name = 'SimHei'
    
    def _process_paragraph(self, elem):
        """处理段落元素。"""
        if not elem.get_text().strip():
            # 空段落
            self.doc.add_paragraph()
            return
        
        p = self.doc.add_paragraph()
        
        # 检查段落中是否有公式标签
        text = str(elem)  # 使用HTML字符串而不是get_text()
        formula_pattern = r'<formula>(.*?)</formula>'
        if re.search(formula_pattern, text):
            # 处理带有公式的段落
            parts = []
            last_end = 0
            
            # 将文本分割为公式和非公式部分
            for match in re.finditer(formula_pattern, text):
                # 公式前的文本
                if match.start() > last_end:
                    parts.append(("text", text[last_end:match.start()]))
                
                # 公式
                parts.append(("formula", match.group(1)))
                last_end = match.end()
            
            # 最后一个公式后的文本
            if last_end < len(text):
                parts.append(("text", text[last_end:]))
            
            # 处理每个部分
            for part_type, content in parts:
                if part_type == "text":
                    if content.strip():
                        # 创建临时段落元素来处理这部分文本
                        temp_soup = BeautifulSoup(content, 'html.parser')
                        if temp_soup.find():
                            # 处理子元素以处理格式
                            self._process_html_fragment(p, temp_soup)
                        else:
                            # 只是纯文本
                            clean_text = self._clean_markdown_text(content)
                            if clean_text.strip():
                                run = p.add_run(clean_text.strip())
                                run.font.name = 'SimHei'
                else:  # 公式
                    print(f"处理公式: {content}")
                    self._process_formula(p, content)
        else:
            # 没有公式的普通段落 - 使用现有处理方式
            # 处理子元素以处理格式
            for child in elem.children:
                if child.name is None:  # 文本
                    if child.string:
                        clean_text = self._clean_markdown_text(child.string)
                        run = p.add_run(clean_text)
                        run.font.name = 'SimHei'
                elif child.name == 'strong' or child.name == 'b':  # 粗体
                    clean_text = self._clean_markdown_text(child.get_text())
                    run = p.add_run(clean_text)
                    run.bold = True
                    run.font.name = 'SimHei'
                elif child.name == 'em' or child.name == 'i':  # 斜体
                    clean_text = self._clean_markdown_text(child.get_text())
                    run = p.add_run(clean_text)
                    run.italic = True
                    run.font.name = 'SimHei'
                elif child.name == 'code':  # 内联代码
                    clean_text = self._clean_markdown_text(child.get_text())
                    run = p.add_run(clean_text)
                    run.font.name = 'SimHei'  # 所有文本都使用黑体
                    run.font.size = Pt(9)
                    # 添加浅灰色高亮
                    shading_elm = self._create_element('w:shd')
                    self._create_attribute(shading_elm, 'w:val', 'clear')
                    self._create_attribute(shading_elm, 'w:color', 'auto')
                    self._create_attribute(shading_elm, 'w:fill', 'F0F0F0')  # 非常浅的灰色
                    run._element.rPr.append(shading_elm)
                else:
                    clean_text = self._clean_markdown_text(child.get_text())
                    run = p.add_run(clean_text)
                    run.font.name = 'SimHei'
    
    def _process_list_item(self, elem, style='List Bullet'):
        """处理列表项元素。"""
        p = self.doc.add_paragraph(style=style)
        
        # 处理子元素以处理列表项中的格式
        for child in elem.children:
            if child.name == 'p':  # 列表项文本在段落中
                for sub_child in child.children:
                    if sub_child.name is None:  # 文本
                        if sub_child.string:
                            clean_text = self._clean_markdown_text(sub_child.string)
                            run = p.add_run(clean_text)
                            run.font.name = 'SimHei'
                    elif sub_child.name == 'strong' or sub_child.name == 'b':  # 粗体
                        clean_text = self._clean_markdown_text(sub_child.get_text())
                        run = p.add_run(clean_text)
                        run.bold = True
                        run.font.name = 'SimHei'
                    elif sub_child.name == 'em' or sub_child.name == 'i':  # 斜体
                        clean_text = self._clean_markdown_text(sub_child.get_text())
                        run = p.add_run(clean_text)
                        run.italic = True
                        run.font.name = 'SimHei'
                    else:
                        clean_text = self._clean_markdown_text(sub_child.get_text())
                        run = p.add_run(clean_text)
                        run.font.name = 'SimHei'
            elif child.name == 'ul' or child.name == 'ol':  # 嵌套列表
                # 在当前项目之后处理嵌套列表
                self._process_list(child, level=2)
            else:
                # 只添加文本
                text = child.string if child.string else child.get_text()
                if text:
                    clean_text = self._clean_markdown_text(text)
                    run = p.add_run(clean_text)
                    run.font.name = 'SimHei'
    
    def _process_list(self, elem, level=1):
        """处理无序或有序列表元素。"""
        style = 'List Bullet' if elem.name == 'ul' else 'List Number'
        if level > 1:
            style = f'{style} {level}'
        
        for li in elem.find_all('li', recursive=False):
            self._process_list_item(li, style=style)
    
    def _process_table(self, elem):
        """处理表格元素。"""
        rows = elem.find_all('tr')
        if not rows:
            return
        
        # 从第一行确定列数
        first_row = rows[0]
        cells = first_row.find_all(['th', 'td'])
        col_count = len(cells)
        if col_count == 0:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=col_count)
        table.style = 'Table Grid'
        
        # 处理每一行
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            
            for j, cell in enumerate(cells):
                if j < col_count:  # 确保不超过列数
                    # 添加单元格内容
                    text = cell.get_text().strip()
                    clean_text = self._clean_markdown_text(text)
                    table.cell(i, j).text = clean_text
                    
                    # 使标题单元格加粗并确保所有单元格使用黑体
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'SimHei'
                            if cell.name == 'th' or i == 0:
                                run.bold = True
    
    def _process_code_block(self, elem):
        """处理预格式化代码块。"""
        code_text = elem.get_text()
        self._format_code_block(code_text)
    
    def _html_to_docx(self, html):
        """将HTML转换为Word文档。"""
        # 如果需要，通过包装在适当的标签中修复HTML
        if not html.strip().startswith('<html>'):
            html = f'<html><body>{html}</body></html>'
            
        soup = BeautifulSoup(html, 'html.parser')
        
        # 检查body是否存在，如果不存在则创建一个
        body = soup.body
        if not body:
            print("警告: HTML中未找到body标签")
            return
        
        # 按顺序处理元素
        for elem in body.children:
            if elem.name is None:  # 跳过非标签元素
                continue
                
            elif elem.name == 'h1':
                self._process_heading(elem, 1)
            elif elem.name == 'h2':
                self._process_heading(elem, 2)
            elif elem.name == 'h3':
                self._process_heading(elem, 3)
            elif elem.name == 'h4':
                self._process_heading(elem, 4)
            elif elem.name == 'h5':
                self._process_heading(elem, 5)
            elif elem.name == 'h6':
                self._process_heading(elem, 6)
            elif elem.name == 'p':
                self._process_paragraph(elem)
            elif elem.name == 'ul' or elem.name == 'ol':
                self._process_list(elem)
            elif elem.name == 'table':
                self._process_table(elem)
            elif elem.name == 'pre':
                self._process_code_block(elem)
            elif elem.name == 'hr':
                p = self.doc.add_paragraph('─' * 50)  # 简单的水平线
                for run in p.runs:
                    run.font.name = 'SimHei'
            elif elem.name == 'blockquote':
                # 简单的块引用作为缩进文本
                for p in elem.find_all('p'):
                    paragraph = self.doc.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.right_indent = Inches(0.5)
                    clean_text = self._clean_markdown_text(p.get_text())
                    run = paragraph.add_run(clean_text)
                    run.italic = True
                    run.font.name = 'SimHei'
    
    def _sequential_parse(self, markdown_text):
        """逐行按顺序解析Markdown。"""
        # 预处理以处理多行公式 
        formula_map = {}
        formula_counter = 0
        
        def replace_formula(match):
            nonlocal formula_counter
            formula_content = match.group(1)
            placeholder = f"FORMULA_PLACEHOLDER_{formula_counter}"
            formula_map[placeholder] = formula_content
            formula_counter += 1
            return f"<formula>{placeholder}</formula>"
        
        # 使用re.DOTALL跨行匹配，将所有公式替换为占位符
        processed_text = re.sub(r'<formula>(.*?)</formula>', replace_formula, markdown_text, flags=re.DOTALL)
        
        # 创建新文档进行顺序处理
        self.doc = Document()
        self._set_default_styles()
        
        # 常规处理变量
        in_code_block = False
        code_block_content = []
        code_block_language = ""
        in_table = False
        table_content = []
        in_diagram = False
        diagram_content = []
        
        # 将处理后的Markdown分割为行
        lines = processed_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 检查代码块 - 修改这部分以更好地识别代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    # 代码块结束
                    self._format_code_block('\n'.join(code_block_content), code_block_language)
                    code_block_content = []
                    code_block_language = ""
                    in_code_block = False
                else:
                    # 代码块开始
                    in_code_block = True
                    # 提取指定的语言
                    language_part = line.strip()[3:].strip()
                    code_block_language = language_part
                i += 1
                continue
            
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # 检查ASCII艺术图（带有多个特殊字符的行）
            if ('+' in line and '-' in line and '|' in line) or \
               re.search(r'[+\-|↓→←↑]{3,}', line) and not in_diagram and not in_table:
                in_diagram = True
                diagram_content.append(line)
                i += 1
                continue
            elif in_diagram:
                if line.strip() == '' and len(diagram_content) > 0:
                    # 在空行处结束图表
                    self._format_diagram('\n'.join(diagram_content))
                    diagram_content = []
                    in_diagram = False
                elif not any(c in line for c in '+|-|↓→←↑') and len(diagram_content) > 1:
                    # 在没有图表字符的行处结束图表
                    self._format_diagram('\n'.join(diagram_content))
                    diagram_content = []
                    in_diagram = False
                    # 不增加i，处理这一行
                    continue
                else:
                    diagram_content.append(line)
                    i += 1
                    continue
            
            # 检查表格
            if line.startswith('|') and not in_table:
                # 检查下一行是否包含分隔符
                if i+1 < len(lines) and lines[i+1].startswith('|') and all(c == '-' or c == '|' or c == ':' or c == ' ' for c in lines[i+1]):
                    in_table = True
                    table_content.append(line)
            elif in_table and line.startswith('|'):
                table_content.append(line)
            elif in_table:
                # 表格结束，处理它
                self._process_markdown_table('\n'.join(table_content))
                table_content = []
                in_table = False
                # 这里不增加i，处理这一行
                continue
            
            # 如果在表格或代码块中则跳过
            if in_table or in_code_block or in_diagram:
                i += 1
                continue
            
            # 清理并简化公式处理
            if '<formula>' in line:
                p = self.doc.add_paragraph()
                formula_pattern = r'<formula>(FORMULA_PLACEHOLDER_\d+)</formula>'
                parts = []
                last_end = 0
                
                for match in re.finditer(formula_pattern, line):
                    # 公式前的文本
                    if match.start() > last_end:
                        parts.append(("text", line[last_end:match.start()]))
                    
                    # 从映射中获取实际的公式内容
                    placeholder = match.group(1)
                    formula_content = formula_map[placeholder]
                    parts.append(("formula", formula_content))
                    last_end = match.end()
                
                # 最后一个公式后的文本
                if last_end < len(line):
                    parts.append(("text", line[last_end:]))
                
                # 处理每个部分
                for part_type, content in parts:
                    if part_type == "text":
                        if content.strip():
                            clean_text = self._clean_markdown_text(content)
                            run = p.add_run(clean_text)
                            run.font.name = 'SimHei'
                    else:  # 公式
                        self._process_formula(p, content)
                
                i += 1
                continue
            
            # 清除星号和其他Markdown格式
            clean_line = self._clean_markdown_text(line)
            
            # 检查类似"1.2 关键假设"的章节标题
            section_heading_match = re.match(r'^(\d+\.\d+)\s+(.+)$', clean_line.strip())
            if section_heading_match:
                heading_text = clean_line.strip()
                heading = self.doc.add_heading(heading_text, level=2)
                # 确保标题使用黑体
                for run in heading.runs:
                    run.font.name = 'SimHei'
                i += 1
                continue
            
            # 特殊中文标题模式
            chinese_heading_match = re.match(r'^([一二三四五六七八九十]+)[、\.](.*)$', clean_line.strip())
            if chinese_heading_match:
                heading_text = clean_line.strip()
                heading = self.doc.add_heading(heading_text, level=1)
                # 确保标题使用黑体
                for run in heading.runs:
                    run.font.name = 'SimHei'
                i += 1
                continue
                
            decimal_heading_match = re.match(r'^(\d+\.\d+)\s+(.*)$', clean_line.strip())
            if decimal_heading_match:
                heading_text = clean_line.strip()
                heading = self.doc.add_heading(heading_text, level=2)
                # 确保标题使用黑体
                for run in heading.runs:
                    run.font.name = 'SimHei'
                i += 1
                continue
            
            # 特殊处理中文标题
            if clean_line.startswith('# '):
                heading = self.doc.add_heading(clean_line[2:].strip(), level=1)
                for run in heading.runs:
                    run.font.name = 'SimHei'
            elif clean_line.startswith('## '):
                heading = self.doc.add_heading(clean_line[3:].strip(), level=2)
                for run in heading.runs:
                    run.font.name = 'SimHei'
            elif clean_line.startswith('### '):
                heading = self.doc.add_heading(clean_line[4:].strip(), level=3)
                for run in heading.runs:
                    run.font.name = 'SimHei'
            elif clean_line.startswith('#### '):
                heading = self.doc.add_heading(clean_line[5:].strip(), level=4)
                for run in heading.runs:
                    run.font.name = 'SimHei'
            
            # 处理常规文本行中的独立"H1:"，"H2:"，"H3:"，"H4:"（不是项目符号）
            elif re.match(r'^H\d+:', clean_line.strip()):
                h_part_match = re.match(r'^(H\d+:)(.*)', clean_line.strip())
                if h_part_match:
                    p = self.doc.add_paragraph()
                    content = h_part_match.group(2).strip()  # 只是内容部分，没有H1:, H2:等
                    
                    # 只添加内容（跳过H标签）
                    content_run = p.add_run(content)
                    content_run.font.name = 'SimHei'
                else:
                    p = self.doc.add_paragraph()
                    run = p.add_run(clean_line.strip())
                    run.font.name = 'SimHei'
            
            # 处理带有H1, H2, H3, H4模式的项目符号
            elif clean_line.strip().startswith('- H') and re.match(r'- H\d+:', clean_line.strip()):
                p = self.doc.add_paragraph(style='List Bullet')
                
                # 分割行以获取H部分和内容部分
                h_part_match = re.match(r'- (H\d+:)(.*)', clean_line.strip())
                if h_part_match:
                    content = h_part_match.group(2).strip()  # 只是内容部分，没有H1:, H2:等
                    
                    # 只添加内容（跳过H标签）
                    content_run = p.add_run(content)
                    content_run.font.name = 'SimHei'
                else:
                    # 如果模式匹配失败，处理整行
                    run = p.add_run(clean_line[2:])
                    run.font.name = 'SimHei'
                
            # 处理常规项目符号和编号列表
            elif clean_line.startswith('- '):
                p = self.doc.add_paragraph(style='List Bullet')
                run = p.add_run(clean_line[2:])
                run.font.name = 'SimHei'
            elif clean_line.startswith('  - '):
                p = self.doc.add_paragraph(style='List Bullet 2')
                run = p.add_run(clean_line[4:])
                run.font.name = 'SimHei'
            elif re.match(r'^\d+\. ', clean_line):
                match = re.match(r'^\d+\. ', clean_line)
                p = self.doc.add_paragraph(style='List Number')
                run = p.add_run(clean_line[match.end():])
                run.font.name = 'SimHei'
            elif clean_line.strip() == '':
                self.doc.add_paragraph()
            else:
                # 常规段落
                p = self.doc.add_paragraph()
                run = p.add_run(clean_line)
                run.font.name = 'SimHei'
            
            i += 1
        
        # 处理任何剩余内容
        if in_table and table_content:
            self._process_markdown_table('\n'.join(table_content))
        if in_code_block and code_block_content:
            self._format_code_block('\n'.join(code_block_content), code_block_language)
        if in_diagram and diagram_content:
            self._format_diagram('\n'.join(diagram_content))
    
    def _process_markdown_table(self, table_markdown):
        """直接处理Markdown表格。"""
        lines = table_markdown.strip().split('\n')
        if len(lines) < 3:  # 至少需要标题、分隔符和一行数据
            return
        
        # 解析标题行以确定列数
        header_row = lines[0].strip('|').split('|')
        col_count = len(header_row)
        row_count = len(lines) - 1  # 不计算分隔线
        
        # 创建表格
        table = self.doc.add_table(rows=row_count, cols=col_count)
        table.style = 'Table Grid'
        
        # 简化单元格中的公式处理
        def process_cell_with_formula(cell, cell_text):
            """处理可能包含公式的单元格的助手函数。"""
            if '<formula>' not in cell_text:
                clean_text = self._clean_markdown_text(cell_text.strip())
                cell.text = clean_text
                return
            
            # 清除单元格内容但保留段落
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    paragraph.clear()
                
            # 使用单元格中的第一个段落
            paragraph = cell.paragraphs[0]
            
            # 处理公式和文本部分
            formula_pattern = r'<formula>(.*?)</formula>'
            parts = []
            last_end = 0
            
            for match in re.finditer(formula_pattern, cell_text):
                if match.start() > last_end:
                    parts.append(("text", cell_text[last_end:match.start()]))
                parts.append(("formula", match.group(1)))
                last_end = match.end()
                
            if last_end < len(cell_text):
                parts.append(("text", cell_text[last_end:]))
            
            for part_type, content in parts:
                if part_type == "text" and content.strip():
                    run = paragraph.add_run(self._clean_markdown_text(content.strip()))
                    run.font.name = 'SimHei'
                elif part_type == "formula":
                    self._process_formula(paragraph, content)
        
        # 使用助手函数处理标题单元格
        header_cells = table.rows[0].cells
        for i, cell_text in enumerate(header_row):
            if i < col_count:
                process_cell_with_formula(header_cells[i], cell_text)
                # 使标题加粗
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'SimHei'
        
        # 使用助手函数处理数据单元格
        data_row_index = 2  # 跳过标题和分隔符
        table_row_index = 1  # 从标题行之后开始
        
        while data_row_index < len(lines) and table_row_index < row_count:
            data_cells = lines[data_row_index].strip('|').split('|')
            row_cells = table.rows[table_row_index].cells
            
            for i, cell_text in enumerate(data_cells):
                if i < col_count:
                    process_cell_with_formula(row_cells[i], cell_text)
                    # 确保字体
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'SimHei'
            
            data_row_index += 1
            table_row_index += 1
    
    def _enforce_simhei_font(self):
        """强制文档中所有元素使用黑体字体。"""
        # 对所有段落应用黑体
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'SimHei'
        
        # 对所有表格应用黑体
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'SimHei'
        
        # 对页眉和页脚应用黑体（如果有）
        for section in self.doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for paragraph in header.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'SimHei'
    
    def convert(self, markdown_text, output_file):
        """
        将Markdown文件转换为Word文档(.docx)
        
        参数:
            input_file (str): Markdown文件的路径
            output_file (str): 输出Word文档的路径
        """
        try:
            # 清理Markdown文本中的转义字符和不需要的模式
            markdown_text = re.sub(r'\\([\\`*_{}[\]()#+.!-])', r'\1', markdown_text)
            
            # 规范化代码块格式 - 确保代码块被正确识别
            # 查找并修复可能不标准的代码块格式
            markdown_text = self._normalize_code_blocks(markdown_text)
            
            # 使用顺序解析来保持原始文档顺序
            self._sequential_parse(markdown_text)
            
            # 再次确保所有文本使用黑体
            self._enforce_simhei_font()
            
            # 设置XML语言属性以强制使用字体
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    rPr = run._element.get_or_add_rPr()
                    lang = OxmlElement('w:lang')
                    self._create_attribute(lang, 'w:val', 'zh-CN')
                    self._create_attribute(lang, 'w:eastAsia', 'zh-CN')
                    rPr.append(lang)
                    
                    # 添加直接字体元素
                    rFonts = OxmlElement('w:rFonts')
                    self._create_attribute(rFonts, 'w:ascii', 'SimHei')
                    self._create_attribute(rFonts, 'w:eastAsia', 'SimHei')
                    self._create_attribute(rFonts, 'w:hAnsi', 'SimHei')
                    self._create_attribute(rFonts, 'w:cs', 'SimHei')
                    rPr.append(rFonts)
            
            # 保存文档
            self.doc.save(output_file)
            print(f"成功将文本转换为 {output_file}")
            return True
            
        except Exception as e:
            print(f"转换文件时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _normalize_code_blocks(self, text):
        """规范化代码块的格式，确保它们能被正确识别。"""
        # 确保代码块的开始和结束标记在各自的行上
        text = re.sub(r'([^\n])```', r'\1\n```', text)
        text = re.sub(r'```([^\n])', r'```\n\1', text)
        
        # 处理没有语言标识的代码块
        lines = text.split('\n')
        result = []
        i = 0
        while i < len(lines):
            line = lines[i]
            if line.strip() == '```':
                result.append('```text')  # 为没有语言的代码块添加默认语言
            else:
                result.append(line)
            i += 1
        
        return '\n'.join(result)

    def _process_html_fragment(self, paragraph, soup):
        """处理HTML片段并将其添加到给定段落。"""
        if not soup:
            return
        
        # 处理汤中的所有节点
        for node in soup.contents:
            if node.name is None:  # 文本节点
                if node.string and node.string.strip():
                    clean_text = self._clean_markdown_text(node.string)
                    if clean_text.strip():
                        run = paragraph.add_run(clean_text)
                        run.font.name = 'SimHei'
            elif node.name == 'strong' or node.name == 'b':  # 粗体
                clean_text = self._clean_markdown_text(node.get_text())
                if clean_text.strip():
                    run = paragraph.add_run(clean_text)
                    run.bold = True
                    run.font.name = 'SimHei'
            elif node.name == 'em' or node.name == 'i':  # 斜体
                clean_text = self._clean_markdown_text(node.get_text())
                if clean_text.strip():
                    run = paragraph.add_run(clean_text)
                    run.italic = True
                    run.font.name = 'SimHei'
            elif node.name == 'code':  # 内联代码
                clean_text = self._clean_markdown_text(node.get_text())
                if clean_text.strip():
                    run = paragraph.add_run(clean_text)
                    run.font.name = 'SimHei'
                    run.font.size = Pt(9)
                    # 添加浅灰色高亮
                    shading_elm = self._create_element('w:shd')
                    self._create_attribute(shading_elm, 'w:val', 'clear')
                    self._create_attribute(shading_elm, 'w:color', 'auto')
                    self._create_attribute(shading_elm, 'w:fill', 'F0F0F0')
                    run._element.rPr.append(shading_elm)
            else:
                # 递归处理子节点
                self._process_html_fragment(paragraph, node)
                
if __name__ == "__main__":
    # input_file = "./output/pre_output.txt"
    # output_file = "./output/pre_output.docx"
    input_file = "./output/gathered_tmp_text.txt"
    output_file = "./output/第二章中间步.docx"
    
    # 转换markdown为docx
    converter = MarkdownToDocxConverter()
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    converter.convert(markdown_text, output_file)